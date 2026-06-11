"""
File parsing utilities for extracting text from PDF and DOCX files.
"""

import os
from typing import Optional
import PyPDF2
import pdfplumber
from docx import Document

from backend.config.settings import ALLOWED_EXTENSIONS, MAX_FILE_SIZE_MB


def validate_file(file_name: str, file_size: int) -> dict:
    """
    Validate uploaded file format and size.
    
    Args:
        file_name: Name of the uploaded file.
        file_size: Size of the file in bytes.
    
    Returns:
        dict with 'valid' boolean and 'message' string.
    """
    ext = os.path.splitext(file_name)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        return {"valid": False, "message": f"Invalid file format. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"}
    
    max_bytes = MAX_FILE_SIZE_MB * 1024 * 1024
    if file_size > max_bytes:
        return {"valid": False, "message": f"File too large. Maximum size: {MAX_FILE_SIZE_MB}MB"}
    
    return {"valid": True, "message": "File is valid."}


def extract_text_from_pdf(file_path: str) -> Optional[str]:
    """
    Extract text from a PDF file using pdfplumber (primary) and PyPDF2 (fallback).
    
    Args:
        file_path: Path to the PDF file.
    
    Returns:
        Extracted text or None if extraction fails.
    """
    text = ""
    
    # Primary extraction with pdfplumber
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception:
        text = ""
    
    # Fallback to PyPDF2 if pdfplumber fails
    if not text.strip():
        try:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception:
            return None
    
    return text.strip() if text.strip() else None


def extract_text_from_docx(file_path: str) -> Optional[str]:
    """
    Extract text from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file.
    
    Returns:
        Extracted text or None if extraction fails.
    """
    try:
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        
        text = "\n".join(paragraphs)
        return text if text.strip() else None
    except Exception:
        return None


def extract_text(file_path: str) -> Optional[str]:
    """
    Extract text from a file based on its extension.
    
    Args:
        file_path: Path to the file.
    
    Returns:
        Extracted text or None if extraction fails.
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        return None


def save_uploaded_file(uploaded_file, upload_dir: str) -> Optional[str]:
    """
    Save an uploaded file to disk.
    
    Args:
        uploaded_file: Streamlit uploaded file object.
        upload_dir: Directory to save the file.
    
    Returns:
        Path to saved file or None if save fails.
    """
    try:
        os.makedirs(upload_dir, exist_ok=True)
        file_path = os.path.join(upload_dir, uploaded_file.name)
        
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        return file_path
    except Exception:
        return None
